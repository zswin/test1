# coding=utf-8
__author__ = 'zs'
#pip install python-barcode;
#pip install python-docx
import barcode
from barcode.writer import ImageWriter
from docx import Document
import xlrd
from PIL import Image

workbook = xlrd.open_workbook('c:/temp/splist.xlsx')
print (workbook.sheet_names())
sheet = workbook.sheet_by_index(0)
print (sheet.ncols, sheet.nrows, sheet.name)
#rows = sheet.row_values(1)
#cols = sheet.col_values(1)
#print (rows, cols)

def f_gen_pic(bc):
    if barcode==None:
        return None
    else:
        EN13 = barcode.get_barcode_class('ean13')
        img_bc = EN13(bc, writer = ImageWriter())
        f = 'c:/temp/bc%s'%bc
        img_bc.save(f)
        fpath = f + '.png'
        img = Image.open(fpath)
        img.resize((int(img.width/3), int(img.height/3)), Image.ANTIALIAS).save(fpath)
        return fpath

doc = Document()
doc.add_heading('test')
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '商品名称'
hdr_cells[2].text = '条码'
hdr_cells[3].text = 'IMG'
for i in range(1, sheet.nrows):
    bc = sheet.cell(i, 0).value
    if not str.isdigit(bc):
        continue
    nm = sheet.cell(i, 1).value
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = nm
    row_cells[2].text = bc
    p = row_cells[3].paragraphs[0]
    run = p.add_run()
    fp = f_gen_pic(bc)
    print (fp)
    run.add_picture(fp)
    print ('processing %s %s'%(nm, bc))
doc.save('c:/temp/split.docx')

